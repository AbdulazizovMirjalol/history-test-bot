from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

from database import add_question, clear_questions, connect, init_db


def clean(text: str) -> str:
    return " ".join(str(text or "").split())


def parse_docx(docx_path: str | Path) -> list[dict]:
    doc = Document(str(docx_path))
    if not doc.tables:
        raise ValueError("No table found in DOCX file.")

    table = doc.tables[0]
    questions: list[dict] = []

    # The uploaded file has duplicated cells because of Word table formatting:
    # index 2 = question, index 4 = correct answer, index 6-9 = alternatives.
    for index, row in enumerate(table.rows[1:], start=1):
        cells = [clean(cell.text) for cell in row.cells]
        if len(cells) < 9:
            continue

        question = cells[2]
        correct = cells[4]
        raw_options = [correct] + cells[6:10]

        unique_options: list[str] = []
        for option in raw_options:
            if option and option not in unique_options:
                unique_options.append(option)

        if not question or not correct:
            continue

        if len(unique_options) != 4:
            print(f"SKIP row {index}: expected 4 unique options, got {len(unique_options)} -> {unique_options}")
            continue

        questions.append(
            {
                "source_number": index,
                "question": question,
                "correct_answer": correct,
                "options": unique_options,
            }
        )

    return questions


def import_questions(docx_path: str | Path, db_path: str | Path, reset: bool = True) -> int:
    questions = parse_docx(docx_path)
    conn = connect(db_path)
    init_db(conn)
    if reset:
        clear_questions(conn)

    for item in questions:
        add_question(
            conn,
            source_number=item["source_number"],
            question=item["question"],
            correct_answer=item["correct_answer"],
            options=item["options"],
        )

    conn.close()
    return len(questions)


def main() -> None:
    parser = argparse.ArgumentParser(description="Import test questions from DOCX into SQLite.")
    parser.add_argument("--docx", default="data/history_tests.docx", help="Path to DOCX file")
    parser.add_argument("--db", default="questions.db", help="Path to SQLite database")
    parser.add_argument("--no-reset", action="store_true", help="Do not clear existing questions first")
    args = parser.parse_args()

    count = import_questions(args.docx, args.db, reset=not args.no_reset)
    print(f"Imported {count} questions into {args.db}")


if __name__ == "__main__":
    main()
